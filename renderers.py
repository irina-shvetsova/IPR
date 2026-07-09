"""
Модуль рендера ИПР.

Принимает структурированный JSON от IPRGenerator и собирает два файла:
    - DOCX по фирменному формату (python-docx);
    - ICS-календарь с ключевыми точками контроля (icalendar).

Структура документа отражает правки Потока 1: вводный текст в начале, риски
внутри каждого направления развития (отдельного раздела рисков нет), источник
обучения указан, сроки — по этапам.
"""

from __future__ import annotations

from datetime import datetime, timedelta
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from icalendar import Calendar, Event

# Фирменный акцент — как в дельтовском референсе
_ACCENT = RGBColor(0x7C, 0x3A, 0xED)
_INK = RGBColor(0x1A, 0x1D, 0x23)
_DIM = RGBColor(0x6B, 0x72, 0x80)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def render_docx(data: dict) -> bytes:
    """Собирает DOCX из структурированного ИПР и возвращает байты файла."""
    doc = Document()
    _set_base_style(doc)

    header = data.get("header", {})
    _title(doc, "ИНДИВИДУАЛЬНЫЙ ПЛАН РАЗВИТИЯ")
    _subtitle(doc, header.get("full_name", ""))
    if header.get("role"):
        _muted(doc, header["role"])
    if header.get("period"):
        _kv(doc, "Период", header["period"])
    if header.get("basis"):
        _kv(doc, "Основание", header["basis"])
    _muted(doc, "Конфиденциальный документ", italic=True)

    # Вводный текст (правка 1)
    if data.get("intro"):
        _spacer(doc)
        _body(doc, data["intro"])

    # Раздел 1
    s1 = data.get("section1", {})
    _h1(doc, "1. КОНТЕКСТ И НАЗНАЧЕНИЕ ПЛАНА")
    if s1.get("purpose"):
        _body(doc, s1["purpose"])
    if s1.get("strengths"):
        _h2(doc, "Сильные стороны, на которые опирается план")
        for it in s1["strengths"]:
            _bullet(doc, f"{it.get('name','')} — {it.get('score','')}. {it.get('note','')}")
    if s1.get("growth_zones"):
        _h2(doc, "Зоны роста, на которые направлен план")
        for it in s1["growth_zones"]:
            _bullet(doc, f"{it.get('name','')} — {it.get('score','')}. {it.get('note','')}")
    if s1.get("source_note"):
        _muted(doc, s1["source_note"], italic=True)

    # Раздел 2
    s2 = data.get("section2", {})
    _h1(doc, "2. КОНТЕКСТ РАЗВИТИЯ")
    if s2.get("narrative"):
        _body(doc, s2["narrative"])
    table = s2.get("table", {})
    if table:
        rows = [
            ("Что мотивирует", table.get("motivates", "")),
            ("Фокус плана", table.get("focus", "")),
            ("Ключевой сдвиг", table.get("key_shift", "")),
            ("Вне фокуса", table.get("out_of_focus", "")),
        ]
        _two_col_table(doc, rows)
    if s2.get("source_note"):
        _muted(doc, s2["source_note"], italic=True)

    # Раздел 3
    s3 = data.get("section3", {})
    _h1(doc, "3. ЗОНЫ РОСТА ПО РЕЗУЛЬТАТАМ 360°")
    if s3.get("intro"):
        _body(doc, s3["intro"])
    for i, z in enumerate(s3.get("zones", []), 1):
        _h2(doc, f"3.{i}. {_zone_heading(z)}")
        _body(doc, z.get("text", ""))

    # Раздел 4 — направления развития, риски внутри каждого
    s4 = data.get("section4", {})
    _h1(doc, "4. НАПРАВЛЕНИЯ РАЗВИТИЯ")
    if s4.get("intro"):
        _body(doc, s4["intro"])
    for d in s4.get("directions", []):
        _h2(doc, f"{d.get('num','')} {d.get('title','')}")
        block = [
            ("Компетенция", d.get("competency", "")),
            ("Основание (360°)", d.get("basis_360", "")),
            ("Опора на сильное", d.get("strength_support", "")),
            ("Цель развития", d.get("goal", "")),
            ("На рабочем месте", _join(d.get("workplace", []))),
            ("Развивающие проекты", _join(d.get("projects", []))),
            ("Выход из зоны комфорта", d.get("comfort_zone", "")),
            ("Критерии достижения", _join(d.get("criteria", []))),
        ]
        ras = d.get("risk_and_support", {})
        if ras:
            risk_text = ras.get("risk", "")
            manage = ras.get("how_to_manage", "")
            support = ras.get("psychological_support", "")
            combined = "\n".join(p for p in [risk_text, manage, support] if p)
            block.append(("Риски и поддержка", combined))
        _two_col_table(doc, block)

    # Раздел 5 — обучение
    s5 = data.get("section5", {})
    _h1(doc, "5. ОБУЧЕНИЕ И ИСТОЧНИКИ РАЗВИТИЯ")
    if s5.get("intro"):
        _body(doc, s5["intro"])
    if s5.get("source_note"):
        _muted(doc, s5["source_note"], italic=True)
    if s5.get("books"):
        _h2(doc, "Книги")
        for grp in s5["books"]:
            if grp.get("theme"):
                p = doc.add_paragraph()
                run = p.add_run(grp["theme"])
                run.bold = True
            for b in grp.get("items", []):
                _bullet(doc, f"{b.get('author','')} «{b.get('title','')}» — {b.get('note','')}")
    if s5.get("internal_formats"):
        _h2(doc, "Форматы внутреннего обучения")
        for f in s5["internal_formats"]:
            _bullet(doc, f)
    if s5.get("external_programs"):
        _h2(doc, "Внешние программы (по желанию)")
        for f in s5["external_programs"]:
            _bullet(doc, f)

    # Раздел 6 — точки контроля как квартальные вопросы саморефлексии
    s7 = data.get("section7", {})
    _h1(doc, "6. ТОЧКИ КОНТРОЛЯ — ВОПРОСЫ ДЛЯ САМОРЕФЛЕКСИИ")
    if s7.get("note"):
        _muted(doc, s7["note"], italic=True)
    for q in s7.get("questions", []):
        quarter = q.get("quarter", "")
        intensity = q.get("intensity", "")
        prefix = f"{quarter}" + (f" ({intensity})" if intensity else "")
        _bullet(doc, f"{prefix}. {q.get('question','')}")

    # Раздел 7 — согласование
    if data.get("section8"):
        _h1(doc, "7. СОГЛАСОВАНИЕ")
        _body(doc, data["section8"])
        _spacer(doc)
        _body(doc, "Сотрудник / подпись, дата: ____________________")
        _body(doc, "Руководитель / подпись, дата: ____________________")
        _body(doc, "HR-партнёр / подпись, дата: ____________________")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# ICS
# ---------------------------------------------------------------------------


def render_ics(data: dict, start_date: datetime | None = None) -> bytes:
    """
    Собирает ICS-календарь из квартальных вопросов саморефлексии.

    Раз в квартал в календаре появляется один вопрос. Вопросы идут с нарастающей
    интенсивностью: первый мягкий, последний — жёсткий и конкретный. Это и есть
    точки контроля плана, а не жёсткие дедлайны.
    """
    cal = Calendar()
    cal.add("prodid", "-//IPR MVP//ru")
    cal.add("version", "2.0")

    start = start_date or datetime.now()
    questions = data.get("section7", {}).get("questions", [])

    for idx, q in enumerate(questions):
        event = Event()
        quarter = q.get("quarter", f"Квартал {idx + 1}")
        event.add("summary", f"ИПР · {quarter}: вопрос для саморефлексии")
        intensity = q.get("intensity", "")
        body = q.get("question", "")
        if intensity:
            body = f"[{intensity}] {body}"
        event.add("description", body)
        event_date = (start + timedelta(days=90 * idx)).date()
        event.add("dtstart", event_date)
        event.add("dtend", event_date + timedelta(days=1))
        cal.add_component(event)

    return cal.to_ical()


# ---------------------------------------------------------------------------
# Вспомогательные функции оформления DOCX
# ---------------------------------------------------------------------------


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = _INK


def _title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _INK


def _subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _ACCENT


def _h1(doc, text):
    _spacer(doc)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _ACCENT


def _h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = _INK


def _body(doc, text):
    for chunk in str(text).split("\n"):
        if chunk.strip():
            doc.add_paragraph(chunk.strip())


def _muted(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.italic = italic
    run.font.size = Pt(9.5)
    run.font.color.rgb = _DIM


def _kv(doc, key, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{key}: ")
    run.bold = True
    p.add_run(str(value))


def _bullet(doc, text):
    doc.add_paragraph(str(text), style="List Bullet")


def _spacer(doc):
    doc.add_paragraph()


def _zone_heading(zone: dict) -> str:
    """
    Собирает заголовок зоны роста, не дублируя оценку.

    Модель иногда вписывает оценку прямо в title («Визионерство — 6,3»),
    поэтому добавляем score отдельно только если его там ещё нет.
    """
    title = str(zone.get("title", "")).strip()
    score = str(zone.get("score", "")).strip()
    if not score or score in title:
        return title
    return f"{title} — {score}"


def _join(items) -> str:
    if isinstance(items, list):
        return "\n".join(f"• {i}" for i in items if str(i).strip())
    return str(items)


def _two_col_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        if not str(value).strip():
            continue
        cells = table.add_row().cells
        run = cells[0].paragraphs[0].add_run(key)
        run.bold = True
        for i, line in enumerate(str(value).split("\n")):
            if i == 0:
                cells[1].paragraphs[0].add_run(line)
            elif line.strip():
                cells[1].add_paragraph(line)
    _set_table_widths(table)


def _set_table_widths(table):
    from docx.shared import Inches

    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)
